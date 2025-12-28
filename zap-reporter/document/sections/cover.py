"""
封面頁生成模組
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches


def add_cover_page(doc: Document, data: dict, base_dir: str, company_name: str):
    """
    生成報告封面頁

    Args:
        doc: Word 文檔物件
        data: ZAP 報告數據
        base_dir: 基礎目錄 (用於尋找 logo)
        company_name: 公司名稱
    """
    # 標題
    doc.add_heading(f'{company_name} - 弱點掃描報告', 0)

    # Logo
    logo_path = os.path.join(base_dir, 'logo.png')
    if os.path.exists(logo_path):
        try:
            doc.add_picture(logo_path, width=Inches(2.0))
        except Exception:
            pass

    # 報告資訊
    doc.add_paragraph(f"掃描工具: OWASP ZAP")
    doc.add_paragraph(f"產生日期: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # 掃描目標
    scan_target = data.get('site', [{}])[0].get('@name', 'Unknown Target')
    doc.add_paragraph(f"掃描目標: {scan_target}")

    doc.add_page_break()
