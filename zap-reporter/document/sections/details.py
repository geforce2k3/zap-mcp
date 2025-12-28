"""
弱點詳情頁生成模組
"""
from typing import Dict, Optional
from docx import Document
from docx.shared import Inches, RGBColor

from config.translations import RISK_MAPPING, translate_title
from services.translator import auto_translate
from services.formatter import clean_html, parse_ai_response
from document.renderer import render_markdown
from document.styles import get_risk_color


def _build_ai_solutions_map(ai_data: Optional[dict]) -> Dict[str, str]:
    """建立 AI 解決方案查找表"""
    ai_solutions_map = {}

    if not ai_data or 'solutions' not in ai_data:
        return ai_solutions_map

    raw_solutions = ai_data['solutions']

    # 處理列表格式
    if isinstance(raw_solutions, list):
        new_solutions = {}
        for item in raw_solutions:
            if isinstance(item, dict):
                for k, v in item.items():
                    new_solutions[k] = v
        raw_solutions = new_solutions

    # 建立查找表 (含正規化版本)
    if isinstance(raw_solutions, dict):
        for k, v in raw_solutions.items():
            ai_solutions_map[k] = v
            norm_k = k.lower().strip()
            if norm_k not in ai_solutions_map:
                ai_solutions_map[norm_k] = v

    return ai_solutions_map


def _find_ai_content(
    eng_name: str,
    tw_name: str,
    ai_solutions_map: Dict[str, str]
) -> Optional[str]:
    """查找對應的 AI 建議內容"""
    if eng_name in ai_solutions_map:
        return ai_solutions_map[eng_name]
    elif tw_name and tw_name in ai_solutions_map:
        return ai_solutions_map[tw_name]
    elif eng_name.lower().strip() in ai_solutions_map:
        return ai_solutions_map[eng_name.lower().strip()]
    return None


def _add_detail_row(table, label: str, content: str, is_md: bool = False, color: RGBColor = None):
    """添加詳情表格列"""
    row = table.add_row()
    row.cells[0].text = label
    cell = row.cells[1]

    if is_md:
        render_markdown(cell, content)
    else:
        cell.text = str(content)
        if color and cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = color


def add_details_section(
    doc: Document,
    data: dict,
    ai_data: Optional[dict] = None
):
    """
    生成弱點詳情頁

    Args:
        doc: Word 文檔物件
        data: ZAP 報告數據
        ai_data: AI 分析數據 (可選)
    """
    doc.add_heading('2. 弱點詳情分析', level=1)

    # 建立 AI 解決方案查找表
    ai_solutions_map = _build_ai_solutions_map(ai_data)

    sites = data.get('site', [])

    for site in sites:
        alerts = site.get('alerts', [])

        for alert in alerts:
            eng_name = alert.get('alert', 'Unknown Alert')
            risk_eng = alert.get('riskdesc', 'Info').split(' ')[0]
            desc = clean_html(alert.get('desc', ''))

            # 翻譯
            tw_name = translate_title(eng_name)
            tw_risk = RISK_MAPPING.get(risk_eng, risk_eng)

            # 查找 AI 建議
            ai_content = _find_ai_content(eng_name, tw_name, ai_solutions_map)
            parsed_ai = parse_ai_response(ai_content) if ai_content else None

            # 弱點標題
            doc.add_heading(tw_name, level=2)

            # 詳情表格
            det_table = doc.add_table(rows=0, cols=2)
            det_table.style = 'Table Grid'
            det_table.columns[0].width = Inches(1.5)
            det_table.columns[1].width = Inches(5.0)

            # 弱點原名
            _add_detail_row(det_table, "弱點原名", eng_name)

            # 風險等級
            risk_color = get_risk_color(risk_eng)
            _add_detail_row(det_table, "風險等級", tw_risk, color=risk_color)

            # 弱點描述
            zh_desc = auto_translate(desc)
            _add_detail_row(det_table, "弱點描述", zh_desc)

            # AI 分析或 ZAP 標準建議
            if parsed_ai:
                if parsed_ai.get('explanation'):
                    _add_detail_row(det_table, "弱點分析 (AI)", parsed_ai['explanation'], is_md=True)

                sol_content = parsed_ai.get('solution') or ai_content
                _add_detail_row(det_table, "修復建議 (AI)", sol_content, is_md=True)

                ref_content = parsed_ai.get('reference')
                source_label = "生成式 AI 建議"
            else:
                solution_text = clean_html(alert.get('solution', ''))
                zh_solution = auto_translate(solution_text)
                _add_detail_row(det_table, "修復建議", zh_solution)

                ref_content = clean_html(alert.get('reference', ''))
                source_label = "ZAP 標準建議"

            # 建議來源
            row = det_table.add_row()
            row.cells[0].text = "建議來源"
            cell = row.cells[1]
            p = cell.paragraphs[0]
            run = p.add_run(source_label)
            if parsed_ai:
                run.bold = True
                run.font.color.rgb = RGBColor(0, 112, 192)

            # 參考資料
            if ref_content:
                cell.add_paragraph("")
                p_ref = cell.add_paragraph()
                p_ref.add_run("參考資料:").bold = True
                render_markdown(cell, ref_content)

            doc.add_paragraph("")
