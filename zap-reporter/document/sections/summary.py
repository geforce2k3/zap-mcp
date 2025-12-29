"""
摘要頁生成模組
"""
import os
from typing import Dict, Optional
from docx import Document
from docx.shared import Inches, RGBColor

from document.styles import set_table_header_style
from document.charts import generate_risk_chart
from document.renderer import render_markdown


def _count_risks(data: dict) -> Dict[str, int]:
    """統計各風險等級數量"""
    stats = {"High": 0, "Medium": 0, "Low": 0, "Informational": 0}

    for site in data.get('site', []):
        for alert in site.get('alerts', []):
            risk_desc = alert.get('riskdesc', 'Info').split(' ')[0]
            if risk_desc in stats:
                stats[risk_desc] += 1
            else:
                stats["Informational"] += 1

    return stats


def _add_stats_table(doc: Document, stats: Dict[str, int]):
    """添加統計表格"""
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    # 標頭
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = '風險等級', '數量 (Count)'
    set_table_header_style(hdr[0])
    set_table_header_style(hdr[1])

    # 資料列
    def fill_row(idx, label, count, color):
        c = table.rows[idx].cells
        r1 = c[0].paragraphs[0].add_run(label)
        r2 = c[1].paragraphs[0].add_run(str(count))
        r1.bold = r2.bold = True
        r1.font.color.rgb = r2.font.color.rgb = color

    fill_row(1, "高風險 (High)", stats['High'], RGBColor(255, 0, 0))
    fill_row(2, "中風險 (Medium)", stats['Medium'], RGBColor(255, 165, 0))
    fill_row(3, "低風險 (Low)", stats['Low'], RGBColor(200, 200, 0))
    fill_row(4, "資訊 (Info)", stats['Informational'], RGBColor(0, 0, 255))

def _add_nmap_summary(doc: Document, nmap_data: dict):
    """[New] 添加基礎設施偵察摘要 (Nmap)"""
    doc.add_heading('1.1 基礎設施偵察摘要 (Nmap)', level=2)
    
    if not nmap_data or not nmap_data.get("hosts"):
        doc.add_paragraph("未發現 Nmap 掃描數據。")
        return

    # 1. 主機資訊表
    doc.add_paragraph("主機與作業系統資訊：")
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'IP 位址', '主機名稱', '作業系統 (OS)'
    for cell in hdr: set_table_header_style(cell)

    for host in nmap_data["hosts"]:
        row = table.add_row().cells
        row[0].text = host["ip"]
        row[1].text = host["hostname"] or "N/A"
        row[2].text = host["os"]

    doc.add_paragraph("")

    # 2. 開放端口與 CVE 表
    doc.add_paragraph("開放端口與潛在漏洞 (CVE)：")
    
    # 建立表格
    port_table = doc.add_table(rows=1, cols=3)
    port_table.style = 'Table Grid'
    phdr = port_table.rows[0].cells
    phdr[0].text, phdr[1].text, phdr[2].text = '端口/協定', '服務版本', '潛在漏洞 / 腳本輸出'
    for cell in phdr: set_table_header_style(cell)

    has_vulns = False
    for host in nmap_data["hosts"]:
        for port in host["ports"]:
            row = port_table.add_row().cells
            row[0].text = f"{port['id']}/{port['protocol']}"
            row[1].text = port['service']
            
            # 處理腳本輸出 (CVE)
            script_text = ""
            for script in port['scripts']:
                # 簡單清理輸出，避免太長
                output = script['output'].strip()
                if len(output) > 500: output = output[:500] + "..."
                script_text += f"[{script['id']}]\n{output}\n"
                has_vulns = True
            
            if not script_text:
                script_text = "無檢測到明顯漏洞"
            
            # 這裡可以使用 render_markdown 稍微美化輸出
            row[2].text = script_text

    if has_vulns:
        p = doc.add_paragraph()
        run = p.add_run("警告：偵測到潛在的 CVE 漏洞，請參閱上表腳本輸出。")
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.bold = True

    doc.add_paragraph("")

def add_summary_section(
    doc: Document,
    data: dict,
    base_dir: str,
    ai_data: Optional[dict] = None,
    nmap_data: Optional[dict] = None
):
    """
    生成報告摘要頁

    Args:
        doc: Word 文檔物件
        data: ZAP 報告數據
        base_dir: 基礎目錄
        ai_data: AI 分析數據 (可選)
    """
    doc.add_heading('1. 掃描結果摘要', level=1)

    if nmap_data:
        _add_nmap_summary(doc, nmap_data)
        doc.add_page_break() # 分頁
        
    doc.add_heading('1.2 應用程式弱點摘要 (ZAP)', level=2)
    # 統計風險
    stats = _count_risks(data)
    total_vulns = sum(stats.values())

    doc.add_paragraph(f"本次掃描共發現 {total_vulns} 個潛在弱點。風險分佈如下：")

    # 圓餅圖
    chart_path = os.path.join(base_dir, "risk_chart.png")
    if generate_risk_chart(stats, chart_path):
        try:
            doc.add_picture(chart_path, width=Inches(4.0))
        except Exception:
            doc.add_paragraph("(圖表載入失敗)")

        # 清理暫存圖片
        if os.path.exists(chart_path):
            os.remove(chart_path)

    # 統計表格
    _add_stats_table(doc, stats)
    doc.add_paragraph("")

    # 高風險警告
    if stats['High'] > 0:
        p = doc.add_paragraph()
        run = p.add_run(f"注意：系統存在 {stats['High']} 個高風險弱點，建議立即進行修復！")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_page_break()

    # AI 總結
    if ai_data and ai_data.get('executive_summary'):
        doc.add_heading('生成式 AI 總結', level=2)
        render_markdown(doc, ai_data['executive_summary'])
        doc.add_paragraph("")

    doc.add_page_break()

    return stats
