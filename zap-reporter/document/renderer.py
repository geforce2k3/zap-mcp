"""
Markdown 渲染模組
"""
import re
from docx.shared import Inches, Pt, RGBColor


def render_markdown(container, text: str):
    """
    將 Markdown 內容渲染到 Word 文檔

    支援:
    - 代碼塊 (```)
    - 標題 (##, ###)
    - 列表 (-, *, 1.)
    - 表格 (Markdown 格式)
    - 內聯格式 (**粗體**, `代碼`)

    Args:
        container: Word 文檔容器 (Document 或 Cell)
        text: Markdown 格式的文字
    """
    if not text:
        return

    lines = str(text).split('\n')
    in_code_block = False
    table_buffer = []

    def _render_inline(paragraph, text_content):
        """渲染內聯格式"""
        token_pattern = re.compile(r'(\*\*.*?\*\*)|(`.*?`)')
        parts = token_pattern.split(text_content)

        for part in parts:
            if not part:
                continue
            if part.startswith("`") and part.endswith("`"):
                # 內聯代碼
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.color.rgb = RGBColor(180, 0, 0)
            elif part.startswith("**") and part.endswith("**"):
                # 粗體
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    def _flush_table(buffer):
        """將 Markdown 表格轉換為 Word 表格"""
        if not buffer:
            return

        rows_data = [line.strip().strip('|').split('|') for line in buffer]
        rows_data = [[c.strip() for c in r] for r in rows_data]

        if not rows_data:
            return

        # 檢測標頭行
        headers = None
        body_start = 0
        if len(rows_data) > 1 and all(set(c) <= set('-: ') for c in rows_data[1]):
            headers = rows_data[0]
            body_start = 2

        body_rows = rows_data[body_start:]
        all_rows = ([headers] if headers else []) + body_rows

        if not all_rows:
            return

        max_cols = max(len(r) for r in all_rows)
        table = container.add_table(rows=len(all_rows), cols=max_cols)
        table.style = 'Table Grid'

        curr_idx = 0
        if headers:
            for j, txt in enumerate(headers):
                if j < len(table.rows[curr_idx].cells):
                    p = table.rows[curr_idx].cells[j].paragraphs[0]
                    p.add_run(txt).bold = True
            curr_idx += 1

        for row in body_rows:
            for j, txt in enumerate(row):
                if j < len(table.rows[curr_idx].cells):
                    _render_inline(table.rows[curr_idx].cells[j].paragraphs[0], txt)
            curr_idx += 1

        container.add_paragraph("")

    for line in lines:
        stripped = line.strip()

        # 處理表格
        if not in_code_block and stripped.startswith('|') and stripped.endswith('|'):
            table_buffer.append(stripped)
            continue
        else:
            if table_buffer:
                _flush_table(table_buffer)
                table_buffer = []

        if not stripped:
            continue

        # 代碼塊
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = container.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(80, 80, 80)
            continue

        # 標題
        if stripped.startswith("### ") or stripped.startswith("## "):
            p = container.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(46, 116, 181)
            continue

        # 列表和普通段落
        p = None
        content = stripped

        if stripped.startswith("- ") or stripped.startswith("* "):
            try:
                p = container.add_paragraph(style='List Bullet')
            except Exception:
                p = container.add_paragraph(style='List Paragraph')
            content = stripped[2:]
        elif re.match(r'^\d+\.\s', stripped):
            try:
                p = container.add_paragraph(style='List Number')
            except Exception:
                p = container.add_paragraph(style='List Paragraph')
            content = re.sub(r'^\d+\.\s', '', stripped)

        if p is None:
            p = container.add_paragraph()

        _render_inline(p, content)

    # 處理剩餘的表格
    if table_buffer:
        _flush_table(table_buffer)
