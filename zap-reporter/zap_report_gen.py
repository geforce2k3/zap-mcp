import json
import re
import os
import time
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

# ==========================================
# 0. 字型與環境設定 (修復亂碼關鍵)
# ==========================================

# 嘗試清除 Matplotlib 快取，確保能讀到新安裝的字型
try:
    cache_dir = fm.get_cachedir()
    if os.path.exists(cache_dir):
        shutil.rmtree(cache_dir)
except: pass

# 設定中文字型優先順序
# 容器內通常是 'Noto Sans CJK TC' 或 'Noto Sans CJK SC'
plt.rcParams['font.family'] = ['sans-serif']
plt.rcParams['font.sans-serif'] = ['Noto Sans CJK TC', 'Noto Sans CJK SC', 'WenQuanYi Micro Hei', 'Microsoft JhengHei', 'SimHei', 'sans-serif']
plt.rcParams['axes.unicode_minus'] = False 

# ==========================================
# 1. 翻譯與對照設定
# ==========================================

try:
    from deep_translator import GoogleTranslator
    HAS_TRANSLATOR = True
    translator = GoogleTranslator(source='auto', target='zh-TW')
except ImportError:
    HAS_TRANSLATOR = False
    translator = None
    print("⚠️ 警告: 找不到 deep-translator 模組，將跳過翻譯功能。")

CACHE_FILE = "/app/data/translation_cache.json"
TRANSLATION_CACHE = {}

if os.path.exists(CACHE_FILE):
    try:
        with open(CACHE_FILE, 'r', encoding='utf-8') as f:
            TRANSLATION_CACHE = json.load(f)
    except: pass

def save_cache():
    try:
        with open(CACHE_FILE, 'w', encoding='utf-8') as f:
            json.dump(TRANSLATION_CACHE, f, ensure_ascii=False, indent=2)
    except: pass

def auto_translate(text):
    if not text or len(text) < 2: return text
    if text in TRANSLATION_CACHE: return TRANSLATION_CACHE[text]
    if not HAS_TRANSLATOR: return text
    try:
        result = translator.translate(text[:4500])
        TRANSLATION_CACHE[text] = result
        return result
    except: return text

# [關鍵修復] 完整的風險映射字典
RISK_MAPPING = {
    "High": "高風險 (High)",
    "Medium": "中風險 (Medium)",
    "Low": "低風險 (Low)",
    "Informational": "資訊 (Info)",
    "False Positive": "誤報 (False Positive)"
}

# [關鍵修復] 完整的弱點對照字典 (TERM_MAPPING)
TERM_MAPPING = {
    "Cross Site Scripting (Reflected)": "反射型跨站腳本攻擊 (XSS)",
    "Cross Site Scripting (Persistent)": "儲存型跨站腳本攻擊 (XSS)",
    "Cross Site Scripting (DOM Based)": "DOM 型跨站腳本攻擊 (XSS)",
    "SQL Injection": "SQL 資料隱碼攻擊",
    "Path Traversal": "路徑遍歷漏洞",
    "Remote File Inclusion": "遠端檔案包含 (RFI)",
    "Server Side Include": "伺服器端包含注入 (SSI)",
    "Cross-Site Request Forgery": "跨站請求偽造 (CSRF)",
    "Directory Browsing": "目錄遍歷/目錄瀏覽",
    "Buffer Overflow": "緩衝區溢位",
    "Format String Error": "格式化字串錯誤",
    "Information Disclosure - Debug Error Messages": "資訊洩漏 - 偵錯錯誤訊息",
    "Information Disclosure - Sensitive Information in URL": "資訊洩漏 - URL 包含敏感資訊",
    "Information Disclosure - Suspicious Comments": "資訊洩漏 - 可疑的程式註解",
    "Weak Authentication Method": "身分驗證機制薄弱",
    "Absence of Anti-CSRF Tokens": "缺乏 Anti-CSRF Token",
    "Missing Anti-clickjacking Header": "遺失防點擊劫持標頭 (Clickjacking)",
    "X-Frame-Options Header Not Set": "未設定 X-Frame-Options 標頭",
    "X-Content-Type-Options Header Missing": "遺失 X-Content-Type-Options 標頭",
    "Strict-Transport-Security Header Not Set": "未設定 HSTS 安全傳輸標頭",
    "Cookie No HttpOnly Flag": "Cookie 遺失 HttpOnly 屬性",
    "Cookie Without Secure Flag": "Cookie 遺失 Secure 屬性",
    "Application Error Disclosure": "應用程式錯誤資訊揭露",
    "Private IP Disclosure": "內部 IP 位址洩漏",
    "Session ID in URL Rewrite": "Session ID 暴露於 URL",
    "Source Code Disclosure": "原始碼洩漏",
    "AWS Identity and Access Management (IAM)": "AWS 身分與存取管理",
    "Amazon S3 (Simple Storage Service)": "Amazon S3 物件儲存服務",
    "CloudTrail": "AWS 操作紀錄稽核服務",
    "Cloud IAM": "Google Cloud 身分與存取管理",
}

# ==========================================
# 2. 輔助函式
# ==========================================

def clean_html(raw_html):
    if raw_html is None: return ""
    cleanr = re.compile('<.*?>')
    return re.sub(cleanr, '', raw_html).strip()

def translate_title(english_title):
    # 若找不到對應，則直接回傳原文
    return TERM_MAPPING.get(english_title, english_title)

def set_table_header_style(cell):
    paragraphs = cell.paragraphs
    for p in paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(12)

def generate_risk_chart(stats, output_img_path):
    labels = []
    sizes = []
    colors = []
    mapping = {
        "High": ("高風險", "#ff0000"),
        "Medium": ("中風險", "#ffa500"),
        "Low": ("低風險", "#ffff00"),
        "Informational": ("資訊", "#0000ff")
    }
    for key, (label, color) in mapping.items():
        if stats.get(key, 0) > 0:
            labels.append(f"{label} ({stats[key]})")
            sizes.append(stats[key])
            colors.append(color)
            
    if not sizes: return False

    try:
        plt.figure(figsize=(4, 3))
        plt.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
        plt.axis('equal')
        plt.title("弱點風險分佈") # 這裡會用到中文字型
        plt.tight_layout()
        plt.savefig(output_img_path)
        plt.close()
        return True
    except Exception as e:
        print(f"繪圖失敗: {e}")
        return False

def parse_ai_response(text):
    sections = {'explanation': '', 'solution': '', 'reference': ''}
    current_section = None
    buffer = []
    lines = str(text).split('\n')
    header_regex = re.compile(r'^(#+\s*|\*\*)?(弱點說明|修復建議|解決方法|參考資料|Explanation|Solution|Reference)([:：])?(\*\*)?\s*$')

    for line in lines:
        stripped = line.strip()
        match = header_regex.match(stripped)
        if match:
            if current_section and buffer:
                sections[current_section] = '\n'.join(buffer).strip()
            header_text = match.group(2)
            if '弱點說明' in header_text or 'Explanation' in header_text:
                current_section = 'explanation'
            elif '解決方法' in header_text or '修復建議' in header_text or 'Solution' in header_text:
                current_section = 'solution'
            elif '參考資料' in header_text or 'Reference' in header_text:
                current_section = 'reference'
            buffer = []
            continue
        if current_section:
            buffer.append(line)
        else:
            if stripped and not sections['explanation']:
                current_section = 'explanation'
                buffer.append(line)
    
    if current_section and buffer:
        sections[current_section] = '\n'.join(buffer).strip()
    if not any(sections.values()):
        return {'solution': str(text)}
    return sections

def render_markdown(container, text):
    if not text: return
    lines = str(text).split('\n')
    in_code_block = False
    table_buffer = [] 

    def _render_inline(paragraph, text_content):
        token_pattern = re.compile(r'(\*\*.*?\*\*)|(`.*?`)')
        parts = token_pattern.split(text_content)
        for part in parts:
            if not part: continue
            if part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.color.rgb = RGBColor(180, 0, 0)
            elif part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    def _flush_table(buffer):
        if not buffer: return
        rows_data = [line.strip().strip('|').split('|') for line in buffer]
        rows_data = [[c.strip() for c in r] for r in rows_data]
        if not rows_data: return
        headers = None
        body_start = 0
        if len(rows_data) > 1 and all(set(c) <= set('-: ') for c in rows_data[1]):
            headers = rows_data[0]
            body_start = 2
        body_rows = rows_data[body_start:]
        all_rows = ([headers] if headers else []) + body_rows
        if not all_rows: return
        max_cols = max(len(r) for r in all_rows)
        table = container.add_table(rows=len(all_rows), cols=max_cols)
        table.style = 'Table Grid'
        curr_idx = 0
        if headers:
            for j, txt in enumerate(headers):
                if j < len(table.rows[curr_idx].cells):
                    p = table.rows[curr_idx].cells[j].paragraphs[0]
                    p.add_run(txt).bold = True
            curr_idx += 1
        for row in body_rows:
            for j, txt in enumerate(row):
                if j < len(table.rows[curr_idx].cells):
                    _render_inline(table.rows[curr_idx].cells[j].paragraphs[0], txt)
            curr_idx += 1
        container.add_paragraph("")

    for line in lines:
        stripped = line.strip()
        if not in_code_block and stripped.startswith('|') and stripped.endswith('|'):
            table_buffer.append(stripped)
            continue
        else:
            if table_buffer:
                _flush_table(table_buffer)
                table_buffer = []
        if not stripped: continue
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            p = container.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(80, 80, 80)
            continue
        if stripped.startswith("### ") or stripped.startswith("## "):
            p = container.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(stripped.lstrip("#").strip())
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(46, 116, 181)
            continue
        p = None
        content = stripped
        if stripped.startswith("- ") or stripped.startswith("* "):
            try: p = container.add_paragraph(style='List Bullet')
            except: p = container.add_paragraph(style='List Paragraph')
            content = stripped[2:]
        elif re.match(r'^\d+\.\s', stripped):
            try: p = container.add_paragraph(style='List Number')
            except: p = container.add_paragraph(style='List Paragraph')
            content = re.sub(r'^\d+\.\s', '', stripped)
        if p is None: p = container.add_paragraph()
        _render_inline(p, content)
    if table_buffer: _flush_table(table_buffer)

# ==========================================
# 3. 報告生成主邏輯
# ==========================================

def generate_word_report(json_path, output_path, ai_insights_path=None, company_name="Nextlink MSP"):
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        print(f"讀取 JSON 失敗: {e}")
        return

    ai_data = {}
    if ai_insights_path and os.path.exists(ai_insights_path):
        try:
            with open(ai_insights_path, 'r', encoding='utf-8') as f:
                ai_data = json.load(f)
            print("✅ 成功載入 AI 分析數據！")
        except: pass

    # Fuzzy Lookup Map
    ai_solutions_map = {}
    if 'solutions' in ai_data:
        raw_solutions = ai_data['solutions']
        if isinstance(raw_solutions, list):
            new_solutions = {}
            for item in raw_solutions:
                if isinstance(item, dict):
                    for k, v in item.items():
                        new_solutions[k] = v
            raw_solutions = new_solutions
        if isinstance(raw_solutions, dict):
            for k, v in raw_solutions.items():
                ai_solutions_map[k] = v
                norm_k = k.lower().strip()
                if norm_k not in ai_solutions_map:
                    ai_solutions_map[norm_k] = v

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft JhengHei'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    # 封面
    doc.add_heading(f'{company_name} - 弱點掃描報告', 0)
    base_dir = os.path.dirname(json_path)
    logo_path = os.path.join(base_dir, 'logo.png')
    if os.path.exists(logo_path):
        try: doc.add_picture(logo_path, width=Inches(2.0))
        except: pass
    doc.add_paragraph(f"掃描工具: OWASP ZAP")
    doc.add_paragraph(f"產生日期: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    scan_target = data.get('site', [{}])[0].get('@name', 'Unknown Target')
    doc.add_paragraph(f"掃描目標: {scan_target}")
    doc.add_page_break()

    # 摘要
    doc.add_heading('1. 掃描結果摘要', level=1)
    sites = data.get('site', [])
    stats = {"High": 0, "Medium": 0, "Low": 0, "Informational": 0}
    for site in sites:
        for alert in site.get('alerts', []):
            risk_desc = alert.get('riskdesc', 'Info').split(' ')[0]
            if risk_desc in stats: stats[risk_desc] += 1
            else: stats["Informational"] += 1
    
    total_vulns = sum(stats.values())
    doc.add_paragraph(f"本次掃描共發現 {total_vulns} 個潛在弱點。風險分佈如下：")

    chart_path = os.path.join(base_dir, "risk_chart.png")
    if generate_risk_chart(stats, chart_path):
        try:
            doc.add_picture(chart_path, width=Inches(4.0))
        except:
            doc.add_paragraph("(圖表載入失敗)")
        if os.path.exists(chart_path): os.remove(chart_path)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = '風險等級', '數量 (Count)'
    set_table_header_style(hdr[0])
    set_table_header_style(hdr[1])

    def fill_row(idx, label, count, color):
        c = table.rows[idx].cells
        r1 = c[0].paragraphs[0].add_run(label)
        r2 = c[1].paragraphs[0].add_run(str(count))
        r1.bold = r2.bold = True
        r1.font.color.rgb = r2.font.color.rgb = color

    fill_row(1, "🔴 高風險 (High)", stats['High'], RGBColor(255, 0, 0))
    fill_row(2, "🟠 中風險 (Medium)", stats['Medium'], RGBColor(255, 165, 0))
    fill_row(3, "🟡 低風險 (Low)", stats['Low'], RGBColor(200, 200, 0))
    fill_row(4, "🔵 資訊 (Info)", stats['Informational'], RGBColor(0, 0, 255))
    doc.add_paragraph("")

    if stats['High'] > 0:
        p = doc.add_paragraph()
        run = p.add_run(f"⚠️ 注意：系統存在 {stats['High']} 個高風險弱點，建議立即進行修復！")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
    doc.add_page_break()

    if ai_data.get('executive_summary'):
        doc.add_heading('生成式 AI 總結', level=2)
        render_markdown(doc, ai_data['executive_summary'])
        doc.add_paragraph("")
    doc.add_page_break()

    # 詳情
    doc.add_heading('2. 弱點詳情分析', level=1)

    for site in sites:
        alerts = site.get('alerts', [])
        for alert in alerts:
            eng_name = alert.get('alert', 'Unknown Alert')
            risk_eng = alert.get('riskdesc', 'Info').split(' ')[0]
            desc = clean_html(alert.get('desc', ''))
            
            tw_name = translate_title(eng_name)
            tw_risk = RISK_MAPPING.get(risk_eng, risk_eng)

            ai_content = None
            if eng_name in ai_solutions_map:
                ai_content = ai_solutions_map[eng_name]
            elif tw_name and tw_name in ai_solutions_map:
                ai_content = ai_solutions_map[tw_name]
            elif eng_name.lower().strip() in ai_solutions_map:
                ai_content = ai_solutions_map[eng_name.lower().strip()]
            
            parsed_ai = parse_ai_response(ai_content) if ai_content else None

            doc.add_heading(tw_name, level=2)
            
            det_table = doc.add_table(rows=0, cols=2)
            det_table.style = 'Table Grid'
            det_table.columns[0].width = Inches(1.5)
            det_table.columns[1].width = Inches(5.0)

            def add_row(label, content, is_md=False, color=None):
                row = det_table.add_row()
                row.cells[0].text = label
                cell = row.cells[1]
                if is_md:
                    render_markdown(cell, content)
                else:
                    cell.text = str(content)
                    if color:
                        run = cell.paragraphs[0].runs[0]
                        run.bold = True
                        run.font.color.rgb = color

            add_row("弱點原名", eng_name)
            
            risk_color = None
            if "High" in risk_eng: risk_color = RGBColor(255, 0, 0)
            elif "Medium" in risk_eng: risk_color = RGBColor(255, 165, 0)
            add_row("風險等級", tw_risk, color=risk_color)
            
            zh_desc = auto_translate(desc)
            add_row("弱點描述", zh_desc)

            if parsed_ai:
                if parsed_ai.get('explanation'):
                    add_row("弱點分析 (AI)", parsed_ai['explanation'], is_md=True)
                
                sol_content = parsed_ai.get('solution') or ai_content
                add_row("修復建議 (AI)", sol_content, is_md=True)
                
                ref_content = parsed_ai.get('reference')
                source_label = "生成式 AI 建議"
            else:
                solution_text = clean_html(alert.get('solution', ''))
                zh_solution = auto_translate(solution_text)
                add_row("修復建議", zh_solution)
                ref_content = clean_html(alert.get('reference', ''))
                source_label = "ZAP 標準建議"

            row = det_table.add_row()
            row.cells[0].text = "建議來源"
            cell = row.cells[1]
            p = cell.paragraphs[0]
            run = p.add_run(source_label)
            if parsed_ai:
                run.bold = True
                run.font.color.rgb = RGBColor(0, 112, 192)
            
            if ref_content:
                cell.add_paragraph("") 
                p_ref = cell.add_paragraph()
                p_ref.add_run("參考資料:").bold = True
                render_markdown(cell, ref_content)

            doc.add_paragraph("")

    try:
        doc.save(output_path)
        print(f"報告生成完畢！已儲存至: {output_path}")
        save_cache()
    except Exception as e:
        print(f"儲存失敗: {e}")

if __name__ == "__main__":
    DATA_DIR = "/app/data"
    json_file = os.path.join(DATA_DIR, 'ZAP-Report.json')
    ai_file = os.path.join(DATA_DIR, 'ai_insights.json')
    word_file = os.path.join(DATA_DIR, f'Scan_Report_{datetime.now().strftime("%y%m%d%H%M")}.docx')
    
    if os.path.exists(json_file):
        generate_word_report(json_file, word_file, ai_insights_path=ai_file)
    else:
        print(f"找不到檔案: {json_file}")